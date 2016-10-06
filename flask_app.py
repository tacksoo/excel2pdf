# Started based on this sample:
# http://stackoverflow.com/questions/27628053/uploading-and-downloading-files-with-flask
from flask import Flask, request, redirect, url_for, Response, session
import openpyxl # using version 1.6.1
import docx # using python-docx 0.8.4
import os
import zipstream
import zipfile
import time

app = Flask(__name__)

from schedule import second
app.register_blueprint(second, url_prefix='/schedule')

def get_column_by_name(sheet,name):
  ''' Return Excel column (A,B,C...) based on column heading '''
  for c in range(ord('A'),ord('Z')+1):
    # hard coded to row two, find a better way to do this
    if sheet.cell( chr(c) + '2' ).value == name:
      return chr(c)

def get_names(sheet,column):
  names = ''
  # lastname
  count = 3
  while sheet.cell(column + str(count)).value != None:
    lastname = sheet.cell(column + str(count)).value
    firstname = sheet.cell(chr(ord(column) + 1) + str(count)).value
    if sheet.cell(chr(ord(column) + 8) + str(count)).value == 'PT':
      names += lastname + ',' + firstname + '\n'
    count += 1
  return names

def create_aca_file(firstname,lastname,filepath):
    document = docx.Document(filepath)
    tables = document.tables
    table = tables[0]
    # add name
    cell = table.cell(0,0)
    cell.text = u'To:   ' + firstname + ' ' + lastname
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    cell2 = table.cell(0,1)
    # add date
    cell2.text = u'Date: ' + time.strftime("%m/%d/%Y")
    run2 = cell2.paragraphs[0].runs[0]
    run2.font.bold = True
    document.save('aca_form_' + firstname + '_' + lastname + '.docx')

@app.route('/')
def form():
    return """
        <html>
            <body>
                <h2>ACA documents generator</h2>
                <p>Step 1: Upload directory excel file</p>

                <form action="/transform" method="post" enctype="multipart/form-data">
                    <input type="file" name="data_file" />
                    <input type="submit" />
                </form>

            </body>
        </html>
    """

@app.route('/transform', methods=["POST"])
def transform_view():
    file = request.files['data_file']
    if not file:
        return "No file"

    # openpyxl documentation found here: https://media.readthedocs.org/pdf/openpyxl/latest/openpyxl.pdf
    wb = openpyxl.load_workbook(file)
    sheet = wb.get_sheet_by_name('Last Name')
    lastname_column = get_column_by_name(sheet,'Last Name')
    names = get_names(sheet,lastname_column)
    session['names'] = names
    # if csv list is needed, do the following instead
    #response = make_response(names)
    #response.headers["Content-Disposition"] = "attachment; filename=part_time_faculty.csv"
    return redirect(url_for('get_acaforms')) # url_for calls the function

@app.route('/getacaforms', methods=["GET"])
def get_acaforms():
    my_dir = os.path.dirname(__file__)
    file_path = os.path.join(my_dir, 'aca_form.docx')
    lines = session['names'].split("\n")
    for line in lines:
        if line != "":  # need to check if line is non-empty, ouch!
          tokens = line.split(",")
          create_aca_file(tokens[0],tokens[1],file_path)
    return redirect(url_for('list_files'))


@app.route('/listfiles', methods=['GET'])
def list_files():
    files = os.listdir('.')
    zf = zipfile.ZipFile('aca-forms.zip', mode='w')
    for f in files:
        if f.endswith(".docx"):
            zf.write(f)
    zf.close()
    return """
        <html>
            <body>
                <p><a href='/download'>Click here to download ACA forms without course info</a></p>
                <p>Step 2: Upload schedule.xlsx excel file</p>

                <form action="/schedule" method="post" enctype="multipart/form-data">
                    <input type="file" name="data_file" />
                    <input type="submit" />
                </form>
            </body>
        </html>
    """

@app.route('/download', methods=['GET'])
def download():
    ''' allow user to download zip file '''
    z = zipstream.ZipFile(mode='w')
    #my_dir = os.path.dirname(__file__)
    #file_path = os.path.join(my_dir, 'aca_form.docx')
    z.write('aca-forms.zip')
    response = Response(z, mimetype='application/zip')
    response.headers['Content-Disposition'] = 'attachment; filename={}'.format('aca-forms.zip')
    return response

@app.route('/package', methods=['GET'], endpoint='zipball')
def zipball():
    ''' example of how to zip file and serve to user '''
    z = zipstream.ZipFile(mode='w')
    my_dir = os.path.dirname(__file__)
    file_path = os.path.join(my_dir, 'aca_form.docx')
    z.write(file_path)
    response = Response(z, mimetype='application/zip')
    response.headers['Content-Disposition'] = 'attachment; filename={}'.format('files.zip')
    return response

app.secret_key = os.urandom(24)
